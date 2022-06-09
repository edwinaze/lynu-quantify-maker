import openpyxl
from docx import Document
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment,colors
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import sys

student = {}
addEntrys = {}
subEntrys = {}
global studentamount,year, classes, month,name


def Init():
    wb = openpyxl.load_workbook("./班级名单.xlsx")
    Sheet = wb["Sheet1"]
    maxheight = Sheet.max_row
    global studentamount
    studentamount = maxheight
    names = Sheet["B1" : ("B" + str(maxheight))] # 从头到尾获取名字和学号
    numbers = Sheet["A1" : ("A" + str(maxheight))]
    #print(names)
    for t,v in zip(names,numbers):
        name = t[0].value
        num = v[0].value
        #print(name + " " +  str(num))
        student[name] = {
            "number" : num,
            "total" : 2.0,
            "addentry" :{},
            "subentry" : {},
            "isattend" : True
        }
        #print(student[name])

#   isadd:是否为加分项 entryname:条目名 names[]:作用的学生 notattend:是否为扣全勤条目
def addEntryToStu(isadd, entryname, value, names, notattend):
    for name in names:
        if(isadd):
            student[name]["total"] += value
            student[name]["addentry"][entryname] = value
        else:
            student[name]["total"] -= value
            student[name]["subentry"][entryname] = value
            if (notattend):  # 被扣分的是该扣全勤的条目
                if (student[name]["isattend"]):  # 且还没扣过全勤
                    student[name]["total"] -= 2
                    student[name]["isattend"] = False

def addEntry(isadd, entryname, value, names, department):
    if(isadd):
        if(addEntrys.__contains__(entryname)):
            addEntrys[entryname].append({
                "department" : department,
                "value" : value,
                "names" : names
            })
        else:
            addEntrys[entryname] = [{
                    "department": department,
                    "value": value,
                    "names": names
                }]
    else:
        if (subEntrys.__contains__(entryname)):
            subEntrys[entryname].append({
                "department": department,
                "value": value,
                "names": names
            })
        else:
            subEntrys[entryname] = [{
                "department": department,
                "value": value,
                "names": names
            }]

def mutiAddText():

    # 1.批量文件

    with open("./list.txt","r") as f:
        for line in f.readlines():
            line = line.strip('\n')
            handAddOne(line)

def handAddOne(line):

    # 2. 手动添加单条

    if(len(line) == 0):
        print("格式为： 是否为加分项 部门名 条目名 分值 学生1 学生2 学生3... [是否扣全勤]")
        t = input("请输入：")
        list = t.split()
    else:
        list = line.split()
    if(list[0] == "-1"):
        return False
    if(list[0] == '1'): isadd = True
    if(list[0] == '0'): isadd = False
    department = str(list[1])
    entryname = str(list[2])
    value = float(list[3])
    names = []
    notattend = False
    length = len(list)
    if(list[-1] == '1'):
        notattend = True
        length -= 1
    for i in range(4, length):
        names.append(list[i])
    addEntry(isadd,entryname,value,names,department)
    addEntryToStu(isadd,entryname,value,names,notattend)
    return True

def handAddMuti():

    # 3. 手动添加多条
    line = ""
    while(handAddOne(line)):
        continue


def printExcel():
    # 打印总表
    ## 获取基本信息
    print("请输入基本信息：年级年份 班级 月份 学委名字")
    print("如： 21 软工六 5 何栋宇")
    t = input("请输入：").split()
    year = t[0]
    classes = t[1]
    month = t[2]
    name = t[3]
    ## 创建表格

    wb = openpyxl.Workbook()
    sheet = wb.active

    ## 全表格式

    ### 设置行高和列宽

    for i in range(0, studentamount + 3):
        sheet.row_dimensions[i].height = 18.8
    sheet.column_dimensions['A'].width = 10.85
    sheet.column_dimensions['B'].width = 8.22
    sheet.column_dimensions['C'].width = 63.22
    sheet.column_dimensions['D'].width = 28.89
    sheet.column_dimensions['E'].width = 8.89
    sheet.column_dimensions['F'].width = 11.67

    ### 字体
    #### 表头
    borderSet = Border(
        left=Side(style='thin',color=colors.BLACK),
        right=Side(style='thin',color=colors.BLACK),
        top=Side(style='thin',color=colors.BLACK),
        bottom=Side(style='thin',color=colors.BLACK)
    )
    sheet['A1'].font = Font(name='宋体', size=14)
    sheet['A1'].alignment = Alignment(horizontal = 'center', vertical = 'center')
    sheet['A1'].border = borderSet
    for i in range(1,7):
        cell = sheet.cell(row=2,column=i)
        cell.font = Font(name='仿宋', size = 14)
        cell.alignment = Alignment(horizontal = 'center', vertical = 'center')
        cell.border = borderSet

    #### 表身
    for i in range(3, studentamount+3):
        for j in range(1,7):
            cell = sheet.cell(row=i,column=j)
            cell.font = Font(name="宋体", size = 12)
            cell.border = borderSet
            if(j == 1 or j == 2 or j == 5):
                cell.alignment = Alignment(horizontal='center',vertical='center',)
            else:
                cell.alignment = Alignment(horizontal='left', vertical='center',wrapText=True)

    ## 写入表头
    sheet['A1'].value = "信息技术学院20" + year + "级" + classes + "班 " + month + " 月份德育量化成绩  综评员：" + name +" 辅导员签字：________"
    sheet.merge_cells('A1:F1')

    sheet['A2'].value = "学号"
    sheet['B2'].value = "姓名"
    sheet['C2'].value = "加分项"
    sheet['D2'].value = "减分项"
    sheet['E2'].value = "总分"
    sheet['F2'].value = "学生签名"

    ## 写入数据

    cnt = 3
    for name,s in student.items():
        sheet.cell(row=cnt, column=1).value = s["number"]
        sheet.cell(row=cnt, column=2).value = name

        ### 加分项
        value = ""
        if(s["isattend"]):
            value +="全勤+2"
        else:
            value +="全勤+0"

        for entryname,v in s["addentry"].items():
            if(v % 1 == 0):
                v = int(v)
            value += " " + str(entryname) + "+" + str(v)
        sheet.cell(row=cnt, column=3).value = value

        ### 减分项
        value = ""
        flag = True
        for entryname,v in s["subentry"].items():
            if(v % 1 == 0): v = int(v)
            if(flag):
                value += str(entryname) + "-" + str(v)
                flag = False
            else:
                value += " " + str(entryname) + "-" + str(v)
        sheet.cell(row=cnt, column=4).value = value

        ### 总分
        res = 0
        if(s["total"] % 1 == 0):
            res = int(s["total"])
        else:
            res = s["total"]
        sheet.cell(row=cnt, column=5).value = res
        cnt += 1

    wb.save(year + "级" + classes + "班 " + month + " 月份德育量化总表.xlsx")
    print("制作成功！正在打开结果文件...")
    os.startfile(year + "级" + classes + "班 " + month + " 月份德育量化总表.xlsx")
    os.system("pause")

def printWord():
    print("请输入基本信息：年级年份 班级 月份 学委名字")
    print("如： 21 软工六 5 何栋宇")
    t = input("请输入：").split()
    year = t[0]
    classes = t[1]
    month = t[2]
    name = t[3]
    # 打印附表
    document = Document()
    paragraph = document.add_paragraph()

    run = paragraph.add_run(year + "级" + classes + "班" + month + "月份德育量化附表")

    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(18)
    run.font.bold = True

    paragraph = document.add_paragraph()
    run = paragraph.add_run("一、加分")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(14)
    run.font.bold = True

    cnt = 1
    ## 打印加分
    for entryname, s in addEntrys.items():
        value = entryname + " "
        flag = True
        for v in addEntrys[entryname]:
            number = v["value"]
            if(number % 1 == 0):
                number = int(number)
            if(flag):
                value += "+" + str(number)
                flag = False
            else: value += "、" + "+" +str(number)
        ### 空格打印
        amount = 50 - len(value) - len(addEntrys[entryname][0]["department"])
        value += " "*amount
        ### 加上部门
        value += addEntrys[entryname][0]["department"]
        ### 打印行头
        paragraph = document.add_paragraph()
        run = paragraph.add_run(str(cnt) + "、" + value)
        run.font.name = '楷体'
        run.font.bold = True
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        cnt += 1
        ### 打印内容
        for v in addEntrys[entryname]:
            #### 前面数字
            value = "  "
            number = v["value"]
            if(number % 1 == 0): number = int(number)
            value += "+" + str(number) + " "*6
            paragraph = document.add_paragraph()
            run = paragraph.add_run(value)
            run.font.name = "楷体"
            run.font.bold = True
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            #### 名单
            value = ""
            for n in v["names"]:
                if(len(n) == 2):
                    n = n[0] + "  " + n[1]
                value += n + "  "
            run = paragraph.add_run(value)
            run.font.name = "仿宋"
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    ## 打印减分
    paragraph = document.add_paragraph()
    run = paragraph.add_run("二、减分")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(14)
    run.font.bold = True
    for entryname, s in subEntrys.items():
        value = entryname + " "
        flag = True
        for v in subEntrys[entryname]:
            number = v["value"]
            if(number % 1 == 0): number = int(number)
            if(flag):
                value += "-" + str(number)
                flag = False
            else: value += "、" + "-" +str(number)
        ### 空格打印
        amount = 50 - len(value) - len(subEntrys[entryname][0]["department"])
        value += " "*amount
        ### 加上部门
        value += subEntrys[entryname][0]["department"]
        ### 打印行头
        paragraph = document.add_paragraph()
        run = paragraph.add_run(str(cnt) + "、" + value)
        run.font.name = '楷体'
        run.font.bold = True
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        cnt += 1
        ### 打印内容
        for v in subEntrys[entryname]:
            #### 前面数字
            value = "  "
            number = v["value"]
            if(number % 1 == 0): number = int(v["value"])
            value += "-" + str(number) + " "*6
            paragraph = document.add_paragraph()
            run = paragraph.add_run(value)
            run.font.name = "楷体"
            run.font.bold = True
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            #### 名单
            value = ""
            for n in v["names"]:
                if(len(n) == 2):
                    n = n[0] + "  " + n[1]
                value += n + "  "
            run = paragraph.add_run(value)
            run.font.name = "仿宋"
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    document.save(year + "级" + classes + "班 " + month + " 月份德育量化附表.docx")
    print("\n制作成功！正在打开结果文件...")
    os.startfile(year + "级" + classes + "班 " + month + " 月份德育量化附表.docx")
    os.system("pause")

def menu():
    while (1):
        print("----------量化表自动制作 By Aze--------")
        print("1. 批量文件添加条目（同目录下需有 list.txt）")
        print("2. 手动添加单条条目")
        print("3. 手动批量添加条目（与1不同是在添加完一条后不会结束）")
        print("4. 打印量化表总表")
        print("5. 打印量化表附表")
        print("6. 打印量化表总表 + 打印量化表附表")
        enum = int(input("请输入选择项标号："))
        if (enum == 1):
            mutiAddText()
        if (enum == 2):
            line = ""
            handAddOne(line)
        if (enum == 3):
            handAddMuti()
        if (enum == 4):
            printExcel()
        if (enum == 5):
            printWord()
        if (enum == 6):
            printExcel()
            printWord()
        print("\033c", end="")
        print("")
Init()
print("个人博客：https://www.cnblogs.com/edwinaze/")
print("欢迎传播使用, 请保留著名信息")
print("")
menu()